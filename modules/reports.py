"""
reports.py — Generación de informes Word y HTML para el Sistema de Ausentismo DAP-SSMC
"""
import io
import os
from datetime import datetime
import pandas as pd

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ────────────────────────────────────────────────────────────────────────────
# Paleta de colores institucional
# ────────────────────────────────────────────────────────────────────────────
C = {
    'azul_osc':  '00386B', 'azul_med':  '006FB9', 'azul_clar': 'D6E8F7',
    'gris_clar': 'F2F2F2', 'gris_med':  'D0D0D0',
    'verde':     '1E8B4C', 'amarillo':  'FFC000',
    'naranja':   'FF6600', 'rojo':      'C00000',
    'blanco':    'FFFFFF', 'negro':     '000000',
}
RGB = {k: RGBColor(int(v[:2],16), int(v[2:4],16), int(v[4:],16)) for k,v in C.items()}

SEMAFORO_HEX = {
    'VERDE': C['verde'], 'AMARILLO': C['amarillo'],
    'NARANJA': C['naranja'], 'ROJO': C['rojo'], 'CRITICO': '6A0DAD',
}

# ────────────────────────────────────────────────────────────────────────────
# Helpers Word
# ────────────────────────────────────────────────────────────────────────────
def _shd(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def _border_bottom(paragraph, hex_color: str):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '8')
    b.set(qn('w:space'), '1'); b.set(qn('w:color'), hex_color)
    pBdr.append(b); pPr.append(pBdr)

def _h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text.upper())
    r.bold = True; r.font.size = Pt(13)
    r.font.color.rgb = RGB['azul_osc']; r.font.name = 'Calibri'
    _border_bottom(p, C['azul_med'])
    return p

def _h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(11)
    r.font.color.rgb = RGB['azul_med']; r.font.name = 'Calibri'
    return p

def _h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(10)
    r.font.color.rgb = RGB['azul_osc']; r.font.name = 'Calibri'
    return p

def _body(doc, text, italic=False, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    if indent: p.paragraph_format.left_indent = Cm(indent)
    r = p.add_run(text)
    r.font.size = Pt(10); r.font.name = 'Calibri'
    r.italic = italic
    return p

def _bullet(doc, text, level=1):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(level * 0.8)
    p.paragraph_format.space_before = Pt(1)
    r = p.add_run(text)
    r.font.size = Pt(10); r.font.name = 'Calibri'

def _kpi_box(doc, kpis: list[tuple]):
    """kpis = [(label, value, color_hex), ...]"""
    n = len(kpis)
    tbl = doc.add_table(rows=2, cols=n)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value, color) in enumerate(kpis):
        c_top = tbl.cell(0, i)
        c_bot = tbl.cell(1, i)
        c_top.text = str(value)
        c_bot.text = label
        for cell in [c_top, c_bot]:
            _shd(cell, color)
            for run in cell.paragraphs[0].runs:
                run.font.name = 'Calibri'; run.font.color.rgb = RGB['blanco']
        c_top.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c_bot.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        t_run = c_top.paragraphs[0].runs[0]
        t_run.bold = True; t_run.font.size = Pt(16)
        c_bot.paragraphs[0].runs[0].font.size = Pt(8)
    doc.add_paragraph()

def _table(doc, headers: list, rows_data: list, hdr_color=None, font_size=9):
    tbl = doc.add_table(rows=1+len(rows_data), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_color = hdr_color or C['azul_osc']
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = str(h)
        _shd(cell, hdr_color)
        r = cell.paragraphs[0].runs[0]
        r.bold = True; r.font.size = Pt(font_size)
        r.font.color.rgb = RGB['blanco']; r.font.name = 'Calibri'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows_data):
        bg = C['gris_clar'] if ri % 2 == 0 else C['blanco']
        for ci, val in enumerate(row):
            cell = tbl.rows[ri+1].cells[ci]
            cell.text = str(val)
            _shd(cell, bg)
            r = cell.paragraphs[0].runs[0]
            r.font.size = Pt(font_size); r.font.name = 'Calibri'
    doc.add_paragraph()
    return tbl

def _info_box(doc, text: str, bg_color=None):
    bg = bg_color or C['azul_clar']
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0,0)
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.font.size = Pt(9.5); r.font.name = 'Calibri'
    _shd(cell, bg)
    doc.add_paragraph()


# ────────────────────────────────────────────────────────────────────────────
# GENERADOR DE INFORME WORD
# ────────────────────────────────────────────────────────────────────────────
def generar_word(resumen: dict, periodo_desc: str = '', dotacion_prom: float = 1602) -> bytes:
    """
    Genera el informe Word completo en memoria y retorna los bytes.
    """
    g     = resumen.get('globales', {})
    serie = resumen.get('serie_mensual', pd.DataFrame())
    cesfam = resumen.get('por_cesfam', pd.DataFrame())
    planta = resumen.get('por_planta', pd.DataFrame())
    func   = resumen.get('por_funcionario', pd.DataFrame())
    tipo   = resumen.get('tipo_lm', pd.DataFrame())
    dur    = resumen.get('duracion', pd.DataFrame())
    gest   = resumen.get('df', pd.DataFrame())

    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.59); sec.page_height = Cm(27.94)
    for m in ['left_margin','right_margin','top_margin','bottom_margin']:
        setattr(sec, m, Cm(2.5))

    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'; normal.font.size = Pt(10)

    hoy = datetime.now().strftime('%d de %B de %Y')

    # ── PORTADA ──────────────────────────────────────────────────────────────
    for _ in range(3): doc.add_paragraph()
    _centro(doc, 'SERVICIO DE SALUD METROPOLITANO CENTRAL', 13, True, 'azul_osc')
    _centro(doc, 'DIRECCIÓN DE ATENCIÓN PRIMARIA', 12, True, 'azul_med')
    doc.add_paragraph()

    _tira(doc)  # barra decorativa

    _centro(doc, 'INFORME DE SEGUIMIENTO Y GESTIÓN', 18, True, 'azul_osc')
    _centro(doc, 'DEL AUSENTISMO LABORAL POR LICENCIAS MÉDICAS', 16, True, 'azul_osc')
    doc.add_paragraph()
    _centro(doc, periodo_desc or f'Período analizado — Actualizado al {hoy}', 11, False, 'azul_med')
    doc.add_paragraph()

    _kpi_box(doc, [
        ('Total LM', f"{g.get('n_lm',0):,}", C['azul_osc']),
        ('Días Ausentismo', f"{g.get('dias_total',0):,}", C['azul_med']),
        ('Costo Total', f"${g.get('costo_total',0)/1e6:.1f}M", C['rojo']),
        ('Funcionarios', f"{g.get('n_func',0):,}", C['verde']),
    ])

    _centro(doc, 'Departamento de Calidad de Vida Laboral', 10, False, 'azul_osc')
    _centro(doc, hoy, 10, False, 'azul_med')
    doc.add_page_break()

    # ── I. RESUMEN EJECUTIVO ─────────────────────────────────────────────────
    _h1(doc, 'I. Resumen Ejecutivo')

    ia  = g.get('ia_acumulado', 0)
    ia_m = g.get('ia_mensual', 0)
    tg  = g.get('tasa_gravedad', 0)
    tf  = g.get('tasa_frecuencia', 0)
    meta_m = 2.17

    semaforo_ia = '🟢 dentro de meta' if ia_m <= meta_m else ('🟡 alerta' if ia_m <= 2.60 else '🔴 crítico')

    _info_box(doc,
        f"Índice de Ausentismo acumulado: {ia:.2f} días/funcionario\n"
        f"Índice mensual promedio: {ia_m:.3f} — {semaforo_ia} (meta ≤ {meta_m})\n"
        f"Tasa de Frecuencia: {tf:.2f} LM/100 funcionarios/mes\n"
        f"Tasa de Gravedad: {tg:.2f} días promedio por LM\n"
        f"Total licencias: {g.get('n_lm',0):,} LM en {g.get('n_meses',0)} meses — "
        f"Costo directo: ${g.get('costo_total',0):,.0f}",
        C['azul_clar']
    )

    _h2(doc, '1.1 Indicadores Clave del Período')
    _table(doc,
        ['Indicador', 'Valor', 'Meta', 'Estado'],
        [
            ['Índice Ausentismo mensual',    f"{ia_m:.3f}", '≤ 2,170', '✅' if ia_m<=2.17 else '⚠️' if ia_m<=2.60 else '🔴'],
            ['Tasa Frecuencia (LM/100 func)', f"{tf:.2f}", 'Referencia', '—'],
            ['Tasa Gravedad (días/LM)',       f"{tg:.2f}", 'Referencia', '—'],
            ['% LM corta (≤15 días)',         f"{g.get('pct_corto',0):.1f}%", '< 75%', '—'],
            ['% LM muy corta (≤3 días)',      f"{g.get('pct_muy_corto',0):.1f}%", 'Monitoreo', '—'],
            ['% LM prolongada (≥180 días)',   f"{g.get('pct_prolongado',0):.1f}%", '< 1%', '—'],
            ['Costo directo total LM',        f"${g.get('costo_total',0):,.0f}", '—', '—'],
            ['Costo promedio por día LM',     f"${g.get('costo_por_dia',0):,.0f}", '—', '—'],
            ['LM promedio por funcionario',   str(g.get('lm_por_func',0)), '—', '—'],
            ['Días promedio por funcionario', str(g.get('dias_por_func',0)), '—', '—'],
        ]
    )

    # ── II. EVOLUCIÓN TEMPORAL ───────────────────────────────────────────────
    _h1(doc, 'II. Evolución Temporal del Índice de Ausentismo')
    _h2(doc, '2.1 Serie Mensual')
    if not serie.empty:
        filas = []
        for _, r in serie.iterrows():
            sem = {'VERDE': '🟢', 'AMARILLO': '🟡', 'ROJO': '🔴'}.get(r.get('semaforo',''), '⚪')
            filas.append([
                str(r.get('mes_key','')),
                f"{r.get('n_lm',0):,}",
                f"{r.get('dias',0):,}",
                f"{r.get('ia_mensual',0):.3f}",
                f"{r.get('ia_acumulado',0):.3f}",
                f"{r.get('tg',0):.2f}",
                f"${r.get('costo',0)/1e6:.2f}M",
                sem,
            ])
        _table(doc,
            ['Mes','N° LM','Días','IA Mensual','IA Acumulado','Tasa Gravedad','Costo','Semáforo'],
            filas, font_size=8
        )

    # ── III. POR CESFAM ──────────────────────────────────────────────────────
    _h1(doc, 'III. Análisis por Establecimiento (CESFAM)')
    if not cesfam.empty:
        filas = []
        for _, r in cesfam.iterrows():
            sem_color = SEMAFORO_HEX.get(r.get('semaforo','VERDE'), C['verde'])
            sem_icon = {'VERDE':'🟢','AMARILLO':'🟡','ROJO':'🔴'}.get(r.get('semaforo',''),'⚪')
            filas.append([
                str(r.get('cesfam','')),
                f"{r.get('n_lm',0):,}",
                f"{r.get('dias_total',0):,}",
                f"{r.get('n_funcionarios',0)}",
                f"{r.get('ia_acumulado',0):.2f}",
                f"{r.get('ia_mensual',0):.3f}",
                f"{r.get('tasa_frecuencia',0):.2f}",
                f"{r.get('tasa_gravedad',0):.2f}",
                f"{r.get('pct_corto',0):.1f}%",
                f"{r.get('pct_prolongado',0):.1f}%",
                f"{r.get('n_recurrentes_5',0)}",
                f"${r.get('costo_total',0)/1e6:.2f}M",
                sem_icon,
            ])
        _table(doc,
            ['Establecimiento','N° LM','Días','Func.','IA Acum.','IA Mes',
             'TF','TG','% Corta','% Prolong.','Recur. ≥5','Costo','Sem.'],
            filas, font_size=8
        )

    # ── IV. POR ESTAMENTO ────────────────────────────────────────────────────
    _h1(doc, 'IV. Análisis por Estamento / Planta')
    if not planta.empty:
        filas = [[
            str(r.get('planta','')),
            f"{r.get('n_lm',0):,}",
            f"{r.get('dias_total',0):,}",
            f"{r.get('n_funcionarios',0)}",
            f"{r.get('pct_del_total',0):.1f}%",
            f"{r.get('tasa_gravedad',0):.2f}",
            f"{r.get('lm_por_func',0):.1f}",
            f"{r.get('dias_por_func',0):.1f}",
        ] for _, r in planta.iterrows()]
        _table(doc,
            ['Planta','N° LM','Días','Funcionarios','% del Total',
             'Tasa Gravedad','LM/Func','Días/Func'],
            filas
        )

    # ── V. TIPO DE LM ────────────────────────────────────────────────────────
    _h1(doc, 'V. Distribución por Tipo de Licencia Médica')
    if not tipo.empty:
        filas = [[
            str(r.get('tipo_lm','')),
            f"{r.get('n_lm',0):,}",
            f"{r.get('pct_lm',0):.1f}%",
            f"{r.get('dias',0):,}",
            f"{r.get('pct_dias',0):.1f}%",
            f"${r.get('costo',0)/1e6:.2f}M",
        ] for _, r in tipo.iterrows()]
        _table(doc,
            ['Tipo de LM','N° LM','% LM','Días','% Días','Costo'],
            filas
        )

    # ── VI. TOP FUNCIONARIOS ─────────────────────────────────────────────────
    _h1(doc, 'VI. Funcionarios Prioritarios por Ausentismo')
    _info_box(doc,
        '⚠️  Este listado es de uso RESTRINGIDO — contiene datos sensibles de funcionarios.\n'
        'Uso exclusivo del Departamento de Calidad de Vida Laboral y Dirección DAP.',
        C['amarillo']
    )
    if not func.empty:
        top = func.head(30)
        filas = []
        for _, r in top.iterrows():
            sem = {'VERDE':'🟢','AMARILLO':'🟡','NARANJA':'🟠','ROJO':'🔴'}.get(r.get('semaforo',''),'⚪')
            filas.append([
                str(r.get('rut','')),
                str(r.get('nombre','')),
                str(r.get('unidad','')),
                str(r.get('planta','')),
                f"{r.get('n_lm_total',0)}",
                f"{r.get('dias_total',0)}",
                f"${r.get('costo_total',0):,.0f}",
                str(r.get('proceso','SIN ACCION'))[:20],
                sem,
            ])
        _table(doc,
            ['RUT','Nombre','Unidad','Planta','N° LM','Días','Costo','Proceso','Sem.'],
            filas, font_size=8
        )

    # ── VII. RECOMENDACIONES ──────────────────────────────────────────────────
    _h1(doc, 'VII. Recomendaciones Estratégicas')

    _h2(doc, '7.1 Acciones Inmediatas')
    recs = _generar_recomendaciones(g, cesfam, func)
    for rec in recs['inmediatas']:
        _bullet(doc, rec)

    _h2(doc, '7.2 Acciones de Mediano Plazo')
    for rec in recs['mediano']:
        _bullet(doc, rec)

    _h2(doc, '7.3 Alertas Críticas')
    _info_box(doc, '\n'.join(recs['alertas']), C['rojo'])

    # ── Pie ──────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    _body(doc,
        f'Informe generado automáticamente el {hoy} — '
        'Sistema de Seguimiento de Ausentismo DAP-SSMC — '
        'Departamento Calidad de Vida Laboral.',
        italic=True
    )

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ────────────────────────────────────────────────────────────────────────────
# Helpers portada
# ────────────────────────────────────────────────────────────────────────────
def _centro(doc, text, size, bold, color_key):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold
    r.font.color.rgb = RGB[color_key]; r.font.name = 'Calibri'
    return p

def _tira(doc):
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0,0)
    cell.text = ''
    _shd(cell, C['azul_osc'])
    doc.add_paragraph()


# ────────────────────────────────────────────────────────────────────────────
# Recomendaciones automáticas
# ────────────────────────────────────────────────────────────────────────────
def _generar_recomendaciones(globales: dict, cesfam: pd.DataFrame, func: pd.DataFrame) -> dict:
    inmediatas, mediano, alertas = [], [], []

    ia_m = globales.get('ia_mensual', 0)
    if ia_m > 2.17:
        alertas.append(f'⚠️  IA mensual ({ia_m:.3f}) supera la meta de 2,170 — activar protocolo de intervención.')
    if ia_m > 2.60:
        alertas.append('🔴 IA en nivel CRÍTICO (>2,60) — informe urgente a Dirección SSMC.')

    pct_p = globales.get('pct_prolongado', 0)
    if pct_p > 0.5:
        alertas.append(f'⚠️  {pct_p:.1f}% de LM con duración ≥180 días — revisar casos para envío a COMPIN.')

    n_rec = 0
    if not func.empty and 'n_lm_total' in func.columns:
        n_rec = (func['n_lm_total'] >= 5).sum()
    if n_rec > 0:
        inmediatas.append(f'Activar Programa de Acompañamiento para {n_rec} funcionarios con ≥5 LM en el período.')

    if not cesfam.empty and 'semaforo' in cesfam.columns:
        rojos = cesfam[cesfam['semaforo'] == 'ROJO']['cesfam'].tolist()
        if rojos:
            alertas.append(f'🔴 Establecimientos en nivel ROJO: {", ".join(rojos[:5])} — requieren intervención inmediata.')
        amarillos = cesfam[cesfam['semaforo'] == 'AMARILLO']['cesfam'].tolist()
        if amarillos:
            inmediatas.append(f'Monitoreo intensificado para: {", ".join(amarillos[:5])}.')

    pct_mc = globales.get('pct_muy_corto', 0)
    if pct_mc > 50:
        mediano.append(f'{pct_mc:.1f}% de LM son de 1-3 días — evaluar programa de bienestar preventivo para reducir ausentismo de corto plazo.')

    mediano += [
        'Actualizar dotación mensual para mejorar precisión del denominador del IA.',
        'Implementar análisis de correlación ausentismo-estamento-antigüedad.',
        'Desarrollar modelo predictivo con promedio móvil 3 meses para anticipar cortes.',
    ]

    if not alertas:
        alertas = ['✅ No se detectan alertas críticas en el período analizado.']

    return {'inmediatas': inmediatas or ['Sin acciones inmediatas requeridas.'],
            'mediano': mediano, 'alertas': alertas}


# ────────────────────────────────────────────────────────────────────────────
# GENERADOR DE INFORME HTML
# ────────────────────────────────────────────────────────────────────────────
def generar_html(resumen: dict, periodo_desc: str = '') -> str:
    g       = resumen.get('globales', {})
    serie   = resumen.get('serie_mensual', pd.DataFrame())
    cesfam  = resumen.get('por_cesfam', pd.DataFrame())
    planta  = resumen.get('por_planta', pd.DataFrame())
    tipo    = resumen.get('tipo_lm', pd.DataFrame())
    func    = resumen.get('por_funcionario', pd.DataFrame())

    hoy = datetime.now().strftime('%d/%m/%Y %H:%M')

    def df_to_html(df: pd.DataFrame, cols: list = None) -> str:
        if df.empty:
            return '<p><em>Sin datos para mostrar.</em></p>'
        sub = df[cols] if cols else df
        return sub.to_html(index=False, classes='data-table', border=0,
                           na_rep='—', justify='left')

    sem_color_map = {'VERDE':'#1E8B4C','AMARILLO':'#FFC000','ROJO':'#C00000','NARANJA':'#FF6600'}

    ia_m = g.get('ia_mensual', 0)
    sem_global = 'VERDE' if ia_m <= 2.17 else ('AMARILLO' if ia_m <= 2.60 else 'ROJO')

    # Tabla serie mensual formateada
    serie_html = ''
    if not serie.empty:
        rows_html = ''
        for _, r in serie.iterrows():
            sc = sem_color_map.get(r.get('semaforo','VERDE'), '#1E8B4C')
            rows_html += f"""<tr>
                <td>{r.get('mes_key','')}</td>
                <td>{r.get('n_lm',0):,}</td>
                <td>{r.get('dias',0):,}</td>
                <td><strong>{r.get('ia_mensual',0):.3f}</strong></td>
                <td>{r.get('ia_acumulado',0):.3f}</td>
                <td>{r.get('tg',0):.2f}</td>
                <td>${r.get('costo',0)/1e6:.2f}M</td>
                <td><span style="color:{sc};font-weight:bold">⬤ {r.get('semaforo','')}</span></td>
            </tr>"""
        serie_html = f"""
        <table class="data-table">
          <thead><tr><th>Mes</th><th>N° LM</th><th>Días</th>
          <th>IA Mensual</th><th>IA Acum.</th><th>Tasa Grav.</th>
          <th>Costo</th><th>Semáforo</th></tr></thead>
          <tbody>{rows_html}</tbody>
        </table>"""

    # Tabla CESFAM
    cesfam_html = ''
    if not cesfam.empty:
        rows_html = ''
        for _, r in cesfam.iterrows():
            sc = sem_color_map.get(r.get('semaforo','VERDE'), '#1E8B4C')
            rows_html += f"""<tr>
                <td><strong>{r.get('cesfam','')}</strong></td>
                <td>{r.get('n_lm',0):,}</td>
                <td>{r.get('dias_total',0):,}</td>
                <td>{r.get('n_funcionarios',0)}</td>
                <td><strong>{r.get('ia_acumulado',0):.2f}</strong></td>
                <td>{r.get('ia_mensual',0):.3f}</td>
                <td>{r.get('tasa_frecuencia',0):.2f}</td>
                <td>{r.get('tasa_gravedad',0):.2f}</td>
                <td>{r.get('pct_corto',0):.1f}%</td>
                <td>{r.get('n_recurrentes_5',0)}</td>
                <td>${r.get('costo_total',0)/1e6:.2f}M</td>
                <td><span style="color:{sc};font-weight:bold">⬤</span></td>
            </tr>"""
        cesfam_html = f"""
        <table class="data-table">
          <thead><tr><th>Establecimiento</th><th>N° LM</th><th>Días</th>
          <th>Func.</th><th>IA Acum.</th><th>IA Mes</th><th>TF</th><th>TG</th>
          <th>% Corta</th><th>Recur ≥5</th><th>Costo</th><th>Sem.</th></tr></thead>
          <tbody>{rows_html}</tbody>
        </table>"""

    html = f"""<!DOCTYPE html>
<html lang="es">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>Informe Ausentismo DAP-SSMC</title>
<style>
  :root {{
    --azul-osc: #00386B; --azul-med: #006FB9; --azul-clar: #D6E8F7;
    --verde: #1E8B4C; --rojo: #C00000; --amarillo: #FFC000;
  }}
  * {{ box-sizing: border-box; margin: 0; padding: 0; }}
  body {{ font-family: Calibri, Arial, sans-serif; font-size: 10pt;
          background: #F5F7FA; color: #222; }}
  .header {{ background: var(--azul-osc); color: white; padding: 24px 40px; }}
  .header h1 {{ font-size: 20pt; font-weight: bold; }}
  .header h2 {{ font-size: 13pt; font-weight: normal; margin-top: 4px; color: #B3D4F5; }}
  .header .sub {{ font-size: 9pt; margin-top: 8px; color: #8BB8E8; }}
  .container {{ max-width: 1400px; margin: 0 auto; padding: 24px 32px; }}
  .kpi-grid {{ display: grid; grid-template-columns: repeat(auto-fit, minmax(180px,1fr));
               gap: 16px; margin: 20px 0; }}
  .kpi-card {{ background: white; border-radius: 8px; padding: 20px;
               text-align: center; box-shadow: 0 2px 8px rgba(0,0,0,.08);
               border-top: 4px solid var(--azul-med); }}
  .kpi-card .value {{ font-size: 24pt; font-weight: bold; color: var(--azul-osc); }}
  .kpi-card .label {{ font-size: 9pt; color: #666; margin-top: 4px; }}
  .kpi-card.rojo {{ border-top-color: var(--rojo); }}
  .kpi-card.rojo .value {{ color: var(--rojo); }}
  .kpi-card.verde {{ border-top-color: var(--verde); }}
  .kpi-card.verde .value {{ color: var(--verde); }}
  .section {{ background: white; border-radius: 8px; padding: 24px;
              margin: 16px 0; box-shadow: 0 2px 8px rgba(0,0,0,.06); }}
  .section h2 {{ color: var(--azul-osc); font-size: 13pt; padding-bottom: 8px;
                 border-bottom: 3px solid var(--azul-med); margin-bottom: 16px; }}
  .section h3 {{ color: var(--azul-med); font-size: 11pt; margin: 16px 0 8px; }}
  .data-table {{ width: 100%; border-collapse: collapse; font-size: 9pt; }}
  .data-table th {{ background: var(--azul-osc); color: white; padding: 8px 10px;
                    text-align: left; font-size: 8.5pt; }}
  .data-table td {{ padding: 6px 10px; border-bottom: 1px solid #E8EDF3; }}
  .data-table tr:nth-child(even) {{ background: #F8FAFE; }}
  .data-table tr:hover {{ background: var(--azul-clar); }}
  .alert-box {{ border-radius: 6px; padding: 14px 18px; margin: 10px 0;
                font-size: 9.5pt; line-height: 1.6; }}
  .alert-info  {{ background: var(--azul-clar); border-left: 4px solid var(--azul-med); }}
  .alert-warn  {{ background: #FFF8E1; border-left: 4px solid var(--amarillo); }}
  .alert-danger {{ background: #FDECEA; border-left: 4px solid var(--rojo); }}
  .alert-ok    {{ background: #E8F5E9; border-left: 4px solid var(--verde); }}
  .semaforo-badge {{ display: inline-block; padding: 3px 10px; border-radius: 12px;
                     font-size: 8pt; font-weight: bold; color: white; }}
  .sem-verde   {{ background: var(--verde); }}
  .sem-amarillo {{ background: var(--amarillo); color: #333; }}
  .sem-rojo    {{ background: var(--rojo); }}
  .footer {{ text-align: center; font-size: 8pt; color: #888; padding: 24px;
             border-top: 1px solid #DDD; margin-top: 32px; }}
  @media print {{
    body {{ background: white; }}
    .container {{ padding: 8px; }}
    .section {{ box-shadow: none; border: 1px solid #DDD; }}
  }}
</style>
</head>
<body>

<div class="header">
  <h1>INFORME DE AUSENTISMO LABORAL — LM</h1>
  <h2>Dirección de Atención Primaria · Servicio de Salud Metropolitano Central</h2>
  <div class="sub">{periodo_desc or 'Período de análisis'} · Generado el {hoy}</div>
</div>

<div class="container">

  <!-- KPIs -->
  <div class="kpi-grid">
    <div class="kpi-card"><div class="value">{g.get('n_lm',0):,}</div>
      <div class="label">Total Licencias Médicas</div></div>
    <div class="kpi-card"><div class="value">{g.get('dias_total',0):,}</div>
      <div class="label">Días de Ausentismo</div></div>
    <div class="kpi-card rojo"><div class="value">${g.get('costo_total',0)/1e6:.1f}M</div>
      <div class="label">Costo Directo LM</div></div>
    <div class="kpi-card verde"><div class="value">{g.get('n_func',0):,}</div>
      <div class="label">Funcionarios con LM</div></div>
    <div class="kpi-card"><div class="value">{g.get('ia_acumulado',0):.2f}</div>
      <div class="label">ÍA Acumulado</div></div>
    <div class="kpi-card"><div class="value">{g.get('ia_mensual',0):.3f}</div>
      <div class="label">ÍA Mensual Prom.</div></div>
    <div class="kpi-card"><div class="value">{g.get('tasa_gravedad',0):.1f}d</div>
      <div class="label">Tasa de Gravedad</div></div>
    <div class="kpi-card"><div class="value">{g.get('pct_muy_corto',0):.0f}%</div>
      <div class="label">LM ≤3 días</div></div>
  </div>

  <!-- Serie Mensual -->
  <div class="section">
    <h2>Evolución Mensual del Índice de Ausentismo</h2>
    {serie_html}
  </div>

  <!-- Por CESFAM -->
  <div class="section">
    <h2>Análisis por Establecimiento (CESFAM)</h2>
    {cesfam_html}
  </div>

  <!-- Por Planta -->
  <div class="section">
    <h2>Distribución por Estamento / Planta</h2>
    {df_to_html(planta, ['planta','n_lm','dias_total','n_funcionarios','pct_del_total','tasa_gravedad','lm_por_func','dias_por_func'] if not planta.empty else None)}
  </div>

  <!-- Por Tipo LM -->
  <div class="section">
    <h2>Distribución por Tipo de Licencia Médica</h2>
    {df_to_html(tipo, ['tipo_lm','n_lm','pct_lm','dias','pct_dias','costo'] if not tipo.empty else None)}
  </div>

  <!-- Top Funcionarios -->
  <div class="section">
    <h2>Funcionarios Prioritarios</h2>
    <div class="alert-warn">
      ⚠️ Información de uso restringido — solo Departamento Calidad de Vida Laboral y Dirección DAP.
    </div>
    {df_to_html(func[['rut','nombre','unidad','planta','n_lm_total','dias_total','costo_total','semaforo']].head(30) if not func.empty else func)}
  </div>

</div>

<div class="footer">
  Sistema de Seguimiento de Ausentismo · DAP-SSMC · Departamento Calidad de Vida Laboral · {hoy}
</div>
</body>
</html>"""
    return html
